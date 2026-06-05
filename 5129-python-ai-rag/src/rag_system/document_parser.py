import os
import json
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}


@dataclass
class DocumentElement:
    type: str = "text"
    content: str = ""
    metadata: Dict[str, Any] = None
    page_number: Optional[int] = None
    bounding_box: Optional[Tuple[float, float, float, float]] = None
    text_direction: Optional[str] = None  # 'horizontal' or 'vertical'

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class TableElement(DocumentElement):
    type: str = "table"
    rows: List[List[str]] = None
    headers: List[str] = None

    def __post_init__(self):
        super().__post_init__()
        if self.rows is None:
            self.rows = []
        if self.headers is None:
            self.headers = []


@dataclass
class FigureElement(DocumentElement):
    type: str = "figure"
    caption: str = ""
    image_reference: Optional[str] = None

    def __post_init__(self):
        super().__post_init__()


@dataclass
class ParsedDocument:
    document_id: str
    file_path: str
    file_name: str
    file_type: str
    file_size: int
    content_hash: str
    parsed_at: str
    title: str
    author: Optional[str]
    creation_date: Optional[str]
    total_pages: int
    elements: List[DocumentElement]
    full_text: str
    metadata: Dict[str, Any]

    def to_dict(self) -> Dict[str, Any]:
        return {
            "document_id": self.document_id,
            "file_path": self.file_path,
            "file_name": self.file_name,
            "file_type": self.file_type,
            "file_size": self.file_size,
            "content_hash": self.content_hash,
            "parsed_at": self.parsed_at,
            "title": self.title,
            "author": self.author,
            "creation_date": self.creation_date,
            "total_pages": self.total_pages,
            "elements": [asdict(elem) for elem in self.elements],
            "full_text": self.full_text,
            "metadata": self.metadata,
        }

    def save_to_json(self, output_path: Path) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(self.to_dict(), f, ensure_ascii=False, indent=2)
        logger.info(f"Parsed document saved to: {output_path}")

    @classmethod
    def load_from_json(cls, input_path: Path) -> "ParsedDocument":
        if not input_path.exists():
            raise FileNotFoundError(f"Parsed document file not found: {input_path}")

        with open(input_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        elements = []
        for elem_data in data.get("elements", []):
            elem_type = elem_data.get("type", "text")
            if elem_type == "table":
                elements.append(TableElement(**elem_data))
            elif elem_type == "figure":
                elements.append(FigureElement(**elem_data))
            else:
                elements.append(DocumentElement(**elem_data))

        data["elements"] = elements
        return cls(**data)


def compute_file_hash(file_path: Path, algorithm: str = "sha256") -> str:
    hash_obj = hashlib.new(algorithm)
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hash_obj.update(chunk)
    return hash_obj.hexdigest()


def compute_content_hash(content: str, algorithm: str = "sha256") -> str:
    return hashlib.new(algorithm, content.encode("utf-8")).hexdigest()


class DocumentParser:
    def __init__(self, config=None):
        self.config = config or get_config()
        self.output_dir = self.config.paths.parsed_dir

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. Supported types: {SUPPORTED_EXTENSIONS}"
            )

        logger.info(f"Starting to parse document: {file_path}")

        try:
            if ext == ".pdf":
                parsed_doc = self._parse_pdf(path)
            elif ext in {".docx", ".doc"}:
                parsed_doc = self._parse_word(path)
            elif ext in {".txt"}:
                parsed_doc = self._parse_txt(path)
            elif ext in {".md", ".markdown"}:
                parsed_doc = self._parse_markdown(path)
            else:
                raise ValueError(f"No parser available for file type: {ext}")

            output_file = self.output_dir / f"{parsed_doc.document_id}.json"
            parsed_doc.save_to_json(output_file)

            logger.info(f"Successfully parsed document: {file_path}")
            return parsed_doc

        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {str(e)}", exc_info=True)
            raise RuntimeError(f"Document parsing failed for {file_path}: {str(e)}") from e

    def parse_directory(self, dir_path: str) -> List[ParsedDocument]:
        path = Path(dir_path)
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        documents = []
        for file_path in sorted(path.rglob("*")):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    doc = self.parse(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"Skipping {file_path}: {str(e)}")

        logger.info(f"Parsed {len(documents)} documents from {dir_path}")
        return documents

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        try:
            import pdfplumber
            from pdfminer.layout import LTTextContainer, LTFigure, LTTable
        except ImportError:
            raise ImportError(
                "PDF parsing requires pdfplumber. Install with: pip install pdfplumber pdfminer.six"
            )

        file_size = file_path.stat().st_size
        file_hash = compute_file_hash(file_path)
        doc_id = f"doc_{file_hash[:16]}"

        elements = []
        full_text_parts = []
        total_pages = 0
        title = file_path.stem
        author = None
        creation_date = None

        try:
            with pdfplumber.open(str(file_path)) as pdf:
                if pdf.metadata:
                    title = pdf.metadata.get("Title") or title
                    author = pdf.metadata.get("Author")
                    creation_date = pdf.metadata.get("CreationDate")
                    if creation_date:
                        creation_date = creation_date.replace("D:", "")[:14]

                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    full_text_parts.append(text)

                    text_elements = self._extract_pdf_text_elements(page, page_num)
                    elements.extend(text_elements)

                    table_elements = self._extract_pdf_tables(page, page_num)
                    elements.extend(table_elements)

                    figure_elements = self._extract_pdf_figures(page, page_num)
                    elements.extend(figure_elements)

        except Exception as e:
            logger.warning(f"Advanced PDF parsing failed, falling back to basic extraction: {e}")
            with pdfplumber.open(str(file_path)) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    full_text_parts.append(text)
                    if text.strip():
                        elements.append(
                            DocumentElement(
                                type="text",
                                content=text,
                                metadata={"extraction_method": "basic"},
                                page_number=page_num,
                                text_direction="horizontal",
                            )
                        )

        full_text = "\n\n".join(full_text_parts)
        content_hash = compute_content_hash(full_text)

        return ParsedDocument(
            document_id=doc_id,
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="pdf",
            file_size=file_size,
            content_hash=content_hash,
            parsed_at=datetime.utcnow().isoformat() + "Z",
            title=title,
            author=author,
            creation_date=creation_date,
            total_pages=total_pages,
            elements=elements,
            full_text=full_text,
            metadata={"hash_algorithm": "sha256", "parser": "pdfplumber"},
        )

    def _extract_pdf_text_elements(self, page, page_num: int) -> List[DocumentElement]:
        elements = []
        try:
            text_lines = page.extract_text_lines() or []

            current_paragraph = []
            current_y = None

            for line in text_lines:
                text = line.get("text", "").strip()
                if not text:
                    continue

                y_pos = line.get("bottom", 0)
                direction = self._detect_text_direction(text)

                if current_y is not None and abs(y_pos - current_y) > 20:
                    if current_paragraph:
                        paragraph_text = "\n".join(current_paragraph)
                        elements.append(
                            DocumentElement(
                                type="paragraph",
                                content=paragraph_text,
                                metadata={
                                    "x0": line.get("x0", 0),
                                    "x1": line.get("x1", 0),
                                    "top": line.get("top", 0),
                                    "bottom": line.get("bottom", 0),
                                },
                                page_number=page_num,
                                bounding_box=(
                                    line.get("x0", 0),
                                    line.get("top", 0),
                                    line.get("x1", 0),
                                    line.get("bottom", 0),
                                ),
                                text_direction=direction,
                            )
                        )
                        current_paragraph = []

                current_paragraph.append(text)
                current_y = y_pos

            if current_paragraph:
                paragraph_text = "\n".join(current_paragraph)
                elements.append(
                    DocumentElement(
                        type="paragraph",
                        content=paragraph_text,
                        metadata={},
                        page_number=page_num,
                        text_direction=self._detect_text_direction(paragraph_text),
                    )
                )

        except Exception as e:
            logger.warning(f"Failed to extract text elements from page {page_num}: {e}")
            text = page.extract_text() or ""
            if text.strip():
                elements.append(
                    DocumentElement(
                        type="text",
                        content=text,
                        metadata={"extraction_method": "fallback"},
                        page_number=page_num,
                        text_direction="horizontal",
                    )
                )

        return elements

    def _extract_pdf_tables(self, page, page_num: int) -> List[TableElement]:
        tables = []
        try:
            extracted_tables = page.extract_tables() or []

            for table_idx, table_data in enumerate(extracted_tables):
                if not table_data or len(table_data) < 1:
                    continue

                headers = table_data[0] if len(table_data) > 0 else []
                rows = table_data[1:] if len(table_data) > 1 else []

                rows_str = [[str(cell) if cell else "" for cell in row] for row in rows]
                headers_str = [str(h) if h else "" for h in headers]

                content_parts = [" | ".join(headers_str)] if any(headers_str) else []
                for row in rows_str:
                    content_parts.append(" | ".join(row))
                content = "\n".join(content_parts)

                tables.append(
                    TableElement(
                        type="table",
                        content=content,
                        metadata={
                            "table_index": table_idx,
                            "num_rows": len(rows),
                            "num_columns": len(headers),
                        },
                        page_number=page_num,
                        rows=rows_str,
                        headers=headers_str,
                    )
                )

        except Exception as e:
            logger.warning(f"Failed to extract tables from page {page_num}: {e}")

        return tables

    def _extract_pdf_figures(self, page, page_num: int) -> List[FigureElement]:
        figures = []
        try:
            text = page.extract_text() or ""
            import re

            caption_patterns = [
                r"(?:图|Figure|Fig\.|图表)\s*(\d+[.\d]*)\s*[：:.\-]\s*(.+?)(?=\n\n|\n[A-Z]|$)",
                r"(?:表|Table|Tab\.|表格)\s*(\d+[.\d]*)\s*[：:.\-]\s*(.+?)(?=\n\n|\n[A-Z]|$)",
            ]

            for pattern in caption_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
                for match in matches:
                    fig_num = match.group(1)
                    caption = match.group(2).strip()
                    figure_type = "figure" if "图" in match.group(0) or "fig" in match.group(0).lower() else "table_caption"

                    figures.append(
                        FigureElement(
                            type="figure",
                            content=f"Figure {fig_num}: {caption}",
                            metadata={
                                "figure_number": fig_num,
                                "figure_type": figure_type,
                            },
                            page_number=page_num,
                            caption=caption,
                        )
                    )

        except Exception as e:
            logger.warning(f"Failed to extract figures from page {page_num}: {e}")

        return figures

    def _detect_text_direction(self, text: str) -> str:
        if not text:
            return "horizontal"

        chinese_chars = sum(1 for c in text if "\u4e00" <= c <= "\u9fff")
        total_chars = len([c for c in text if c.strip()])

        if total_chars == 0:
            return "horizontal"

        chinese_ratio = chinese_chars / total_chars

        if chinese_ratio > 0.5 and "\n" in text:
            lines = text.split("\n")
            if len(lines) > 2 and all(len(line.strip()) <= 2 for line in lines[:3]):
                return "vertical"

        return "horizontal"

    def _parse_word(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "Word parsing requires python-docx. Install with: pip install python-docx"
            )

        file_size = file_path.stat().st_size
        file_hash = compute_file_hash(file_path)
        doc_id = f"doc_{file_hash[:16]}"

        elements = []
        full_text_parts = []

        try:
            doc = Document(str(file_path))
            title = doc.core_properties.title or file_path.stem
            author = doc.core_properties.author
            creation_date = (
                doc.core_properties.created.isoformat()
                if doc.core_properties.created
                else None
            )

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name.lower() if para.style.name else ""
                elem_type = "heading" if "heading" in style_name else "paragraph"

                full_text_parts.append(text)
                elements.append(
                    DocumentElement(
                        type=elem_type,
                        content=text,
                        metadata={
                            "style": para.style.name,
                            "level": getattr(para.style, "level", None),
                        },
                        text_direction="horizontal",
                    )
                )

            for table_idx, table in enumerate(doc.tables):
                headers = []
                rows = []

                if len(table.rows) > 0:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                for row in table.rows[1:]:
                    rows.append([cell.text.strip() for cell in row.cells])

                content_parts = [" | ".join(headers)] if any(headers) else []
                for row in rows:
                    content_parts.append(" | ".join(row))
                content = "\n".join(content_parts)

                if content.strip():
                    full_text_parts.append(content)
                    elements.append(
                        TableElement(
                            type="table",
                            content=content,
                            metadata={
                                "table_index": table_idx,
                                "num_rows": len(rows),
                                "num_columns": len(headers),
                            },
                            rows=rows,
                            headers=headers,
                        )
                    )

        except Exception as e:
            logger.warning(f"Advanced Word parsing failed, falling back to text extraction: {e}")
            with open(file_path, "rb") as f:
                raw_text = str(f.read())
            full_text_parts.append(raw_text)
            elements.append(
                DocumentElement(
                    type="text",
                    content=raw_text,
                    metadata={"extraction_method": "fallback"},
                    text_direction="horizontal",
                )
            )

        full_text = "\n\n".join(full_text_parts)
        content_hash = compute_content_hash(full_text)

        return ParsedDocument(
            document_id=doc_id,
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="word",
            file_size=file_size,
            content_hash=content_hash,
            parsed_at=datetime.utcnow().isoformat() + "Z",
            title=title,
            author=author,
            creation_date=creation_date,
            total_pages=1,
            elements=elements,
            full_text=full_text,
            metadata={"hash_algorithm": "sha256", "parser": "python-docx"},
        )

    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        file_size = file_path.stat().st_size
        file_hash = compute_file_hash(file_path)
        doc_id = f"doc_{file_hash[:16]}"

        encodings = ["utf-8", "gbk", "gb18030", "utf-16", "latin-1"]
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise RuntimeError(f"Unable to decode text file: {file_path}")

        full_text = content
        elements = []

        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        for para in paragraphs:
            elements.append(
                DocumentElement(
                    type="paragraph",
                    content=para,
                    metadata={},
                    text_direction=self._detect_text_direction(para),
                )
            )

        content_hash = compute_content_hash(full_text)

        return ParsedDocument(
            document_id=doc_id,
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="text",
            file_size=file_size,
            content_hash=content_hash,
            parsed_at=datetime.utcnow().isoformat() + "Z",
            title=file_path.stem,
            author=None,
            creation_date=None,
            total_pages=1,
            elements=elements,
            full_text=full_text,
            metadata={"hash_algorithm": "sha256", "encoding": used_encoding},
        )

    def _parse_markdown(self, file_path: Path) -> ParsedDocument:
        try:
            import markdown
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "Markdown parsing requires markdown and beautifulsoup4. "
                "Install with: pip install markdown beautifulsoup4"
            )

        file_size = file_path.stat().st_size
        file_hash = compute_file_hash(file_path)
        doc_id = f"doc_{file_hash[:16]}"

        with open(file_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        html = markdown.markdown(
            md_content,
            extensions=["tables", "fenced_code", "toc", "attr_list"],
        )
        soup = BeautifulSoup(html, "html.parser")

        elements = []
        full_text_parts = []

        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "table", "pre"]):
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(element.name[1])
                text = element.get_text().strip()
                if text:
                    full_text_parts.append(text)
                    elements.append(
                        DocumentElement(
                            type="heading",
                            content=text,
                            metadata={"level": level},
                            text_direction=self._detect_text_direction(text),
                        )
                    )

            elif element.name == "p":
                text = element.get_text().strip()
                if text:
                    full_text_parts.append(text)
                    elements.append(
                        DocumentElement(
                            type="paragraph",
                            content=text,
                            metadata={},
                            text_direction=self._detect_text_direction(text),
                        )
                    )

            elif element.name == "table":
                headers = []
                rows = []

                header_row = element.find("tr")
                if header_row:
                    headers = [th.get_text().strip() for th in header_row.find_all(["th", "td"])]

                for row in element.find_all("tr")[1:]:
                    rows.append([td.get_text().strip() for td in row.find_all("td")])

                content_parts = [" | ".join(headers)] if any(headers) else []
                for row in rows:
                    content_parts.append(" | ".join(row))
                content = "\n".join(content_parts)

                if content.strip():
                    full_text_parts.append(content)
                    elements.append(
                        TableElement(
                            type="table",
                            content=content,
                            metadata={
                                "num_rows": len(rows),
                                "num_columns": len(headers),
                            },
                            rows=rows,
                            headers=headers,
                        )
                    )

            elif element.name == "pre":
                text = element.get_text().strip()
                if text:
                    full_text_parts.append(text)
                    elements.append(
                        DocumentElement(
                            type="code",
                            content=text,
                            metadata={},
                            text_direction="horizontal",
                        )
                    )

        full_text = "\n\n".join(full_text_parts)
        content_hash = compute_content_hash(full_text)

        return ParsedDocument(
            document_id=doc_id,
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="markdown",
            file_size=file_size,
            content_hash=content_hash,
            parsed_at=datetime.utcnow().isoformat() + "Z",
            title=file_path.stem,
            author=None,
            creation_date=None,
            total_pages=1,
            elements=elements,
            full_text=full_text,
            metadata={"hash_algorithm": "sha256", "parser": "markdown+bs4"},
        )


from .config import get_config
